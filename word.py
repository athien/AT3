# NORTH METROPOLITAN TAFE
# DIPLOMA IN INFORMATION TECHNOLOGY (ADVANCED NETWORKING)
# ICTPRG443 - APPLY INTERMEDIATE PROGRAMMING

# STUDENT: ALEXANDER THIEN (20093897@tafe.wa.edu.au)
# ASSESSMENT 3: PART B - Word Document Manager
# This is the Python program used to manage *.doc files

# Import the  docx package
import os
from docx import Document
from docx.shared import Inches

# Define file path.
file_path = os.getcwd()

# Create an instance of a Word Document and save it
document = Document()
document.save(file_path+"\\wordwithpython.doc")

# Create headings in the Word Document
def add_headings_with_paragraphs(path, doc):
    """The add_headings_with_paragraphs function adds headings with paragraphs to a Word document. Arguments: file_path and filename"""

    # Add the first heading and paragraph to the Word document
    doc.add_heading("Editing Word documents is fun with Python!", level=0)
    doc.add_paragraph("This document will show you how you can do it, too.")

    # Add the second heading and paragraph to the Word document
    doc.add_heading("Initial steps", level=1)
    doc.add_paragraph("In this chapter you will learn how to:").bold = True
    doc.add_paragraph("Add headings", style="List Bullet")
    doc.add_paragraph("Add paragraphs", style="List Bullet")
    doc.add_paragraph("Insert pictures", style="List Bullet")

    # Add the third heading and paragraph to the Word document
    doc.add_heading("Creating headings", level=1)
    doc.add_paragraph("To create headings, you need to use 'add_heading'. You can also specify the size of the heading using 'level'.")

    # Save the document after applied changes
    doc.save(file_path+"\\wordwithpython.doc")

# Add an image to the Word document
def add_picture(file_path, doc):
    """The add_pictures function will embed an image in a Word document. Arguments: file_path and filename"""

    # Add a heading to the Word document
    doc.add_heading("This is adding a picture in native size to the document using Python!", level=1)

    # Specify the path and dimensions of the image
    doc.add_picture(file_path+"\\zophie.jpg", width=Inches(3), height=Inches(4))

    # Save the document after applied changes
    doc.save(file_path+"\\wordwithpython.doc")

# Get text from a Word document
def get_text(doc):
    """The get_text function will retrieve the contents of a Word file and return the results. Arguments: document"""

    # Define a variable to store the text of the Word document
    full_text = []

    # Iterate through the contents of the Word document and print the contents to the terminal
    for line in doc.paragraphs:
        full_text.append(line.text)
    
    return full_text

# Call the add_headings_and_paragraphs function
add_headings_with_paragraphs(file_path, document)

# Call the add_picture function
add_picture(file_path, document)

# Call the get_text function
get_text(document)
print("The document contains the following text:")
print(get_text(document))
