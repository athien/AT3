# NORTH METROPOLITAN TAFE
# DIPLOMA IN INFORMATION TECHNOLOGY (ADVANCED NETWORKING)
# ICTPRG443 - APPLY INTERMEDIATE PROGRAMMING

# STUDENT: ALEXANDER THIEN (20093897@tafe.wa.edu.au)
# ASSESSMENT 3: PART B - Word Document Manager TEST
# This is the Python program used test the doc.py program

import unittest
from docx import Document
from word import add_headings_with_paragraphs, add_picture, get_text, file_path


class TestWord(unittest.TestCase):
    # Test the add_headings_with_paragraphs function
    def test_add_headings_with_paragraphs(self):
        
        # Create a new Word document
        document = Document()
        
        # Add headings and paragraphs to the Word document
        document.add_heading("Editing Word documents is fun with Python!", level=0)
        document.add_paragraph("This document will show you how you can do it, too.")
        document.add_heading("Initial steps", level=1)
        document.add_paragraph("In this chapter you will learn how to:").bold = True
        document.add_paragraph("Add headings", style="List Bullet")
        document.add_paragraph("Add paragraphs", style="List Bullet")
        document.add_paragraph("Insert pictures", style="List Bullet")
        document.add_heading("Creating headings", level=1)
        document.add_paragraph("To create headings, you need to use 'add_heading'. You can also specify the size of the heading using 'level'.")
        
        # Check Word document for matching headings and paragraphs
        self.assertTrue(document.paragraphs[0].text == "Editing Word documents is fun with Python!")
        self.assertTrue(document.paragraphs[1].text == "This document will show you how you can do it, too.")
        self.assertTrue(document.paragraphs[2].text == "Initial steps")
        self.assertTrue(document.paragraphs[3].text == "In this chapter you will learn how to:")
        self.assertTrue(document.paragraphs[4].text == "Add headings")
        self.assertTrue(document.paragraphs[5].text == "Add paragraphs")
        self.assertTrue(document.paragraphs[6].text == "Insert pictures")
        self.assertTrue(document.paragraphs[7].text == "Creating headings")
        self.assertTrue(document.paragraphs[8].text == "To create headings, you need to use 'add_heading'. You can also specify the size of the heading using 'level'.")
        
    # Test the add_picture function
    def test_add_picture(self):
        
        # Create a new Word document
        document = Document()
        
        # Add a picture to the Word document
        add_picture(file_path, document)
        
        # Check Word document for presence of picture
        self.assertTrue((document.inline_shapes[0]))
        
    # Test the get_text function
    
    def test_get_text(self):
        
        # Create a new Word document
        document = Document()
        
        # Add headings and paragraph to the word document
        document.add_paragraph("Editing Word documents is fun with Python!")
        document.add_paragraph("This document will show you how you can do it, too.")
        document.add_paragraph("Initial steps")
        document.add_paragraph("In this chapter you will learn how to:")
        document.add_paragraph("Add headings")
        document.add_paragraph("Add paragraphs")
        document.add_paragraph("Insert pictures")
        document.add_paragraph("Creating headings")
        document.add_paragraph("To create headings, you need to use 'add_heading'. You can also specify the size of the heading using 'level'.")
        document.add_paragraph("This is adding a picture in native size to the document using Python!")
        document.add_paragraph("")
        
        # Call the get_text function
        full_text = get_text(document)
        
        # Verify the authenticity of the text in the Word document
        self.assertListEqual(full_text, ['Editing Word documents is fun with Python!', 'This document will show you how you can do it, too.', 'Initial steps', 'In this chapter you will learn how to:', 'Add headings', 'Add paragraphs', 'Insert pictures', 'Creating headings', "To create headings, you need to use 'add_heading'. You can also specify the size of the heading using 'level'.", 'This is adding a picture in native size to the document using Python!', ''])

if __name__ == '__main__':
    unittest.main()
